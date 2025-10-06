from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
import PyPDF2
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import black
import io
import re
import json
from typing import Union, List, Dict
from pydantic import BaseModel
from openai import OpenAI
from config import Config

# Validate configuration
Config.validate_config()

# Initialize OpenAI client
try:
    openai_client = OpenAI(api_key=Config.OPENAI_API_KEY)
except Exception as e:
    print(f"Warning: OpenAI client initialization failed: {e}")
    print("Resume parsing will fall back to regex-based parsing")
    openai_client = None

app = FastAPI(title="Resume Text Extractor", version="1.0.0")

# Pydantic models for request/response
class Experience(BaseModel):
    position: str
    company: str
    duration: str
    description: list[str]  # List of bullet points

class Education(BaseModel):
    degree: str
    institution: str
    year: str
    gpa: str = ""
    relevant_coursework: str = ""

class Project(BaseModel):
    name: str
    description: list[str]  # List of bullet points
    technologies: str = ""
    url: str = ""
    duration: str = ""

class Publication(BaseModel):
    title: str
    journal: str = ""
    year: str = ""
    authors: str = ""
    url: str = ""

class Certification(BaseModel):
    name: str
    issuer: str
    year: str = ""
    expiry: str = ""

class VolunteerExperience(BaseModel):
    position: str
    organization: str
    duration: str
    description: list[str]  # List of bullet points

class Reference(BaseModel):
    name: str
    title: str = ""
    company: str = ""
    contact: str = ""

class ResumeData(BaseModel):
    name: str
    email: str
    phone: str
    summary: str
    experience: list[Experience]
    education: list[Education]
    skills: list[str]
    github_profile: str = ""
    linkedin_profile: str = ""
    website: str = ""
    location: str = ""
    publications: list[Publication] = []
    projects: list[Project] = []
    certifications: list[Certification] = []
    languages: list[str] = []
    volunteer_experience: list[VolunteerExperience] = []
    awards: list[str] = []
    references: list[Reference] = []

class ParsedResumeResponse(BaseModel):
    success: bool
    data: ResumeData
    message: str

class ATSAnalysisRequest(BaseModel):
    resume_data: ResumeData
    job_description: str

class ATSScore(BaseModel):
    overall_score: int  # 0-100
    keyword_match_score: int  # 0-100
    experience_relevance: int  # 0-100
    education_fit: int  # 0-100
    skills_alignment: int  # 0-100

class ATSInsight(BaseModel):
    category: str  # "strength", "weakness", "suggestion"
    title: str
    description: str
    impact: str  # "high", "medium", "low"

class ATSRecommendation(BaseModel):
    title: str
    description: str
    priority: str  # "high", "medium", "low"
    effort: str  # "easy", "moderate", "complex"

class ATSAnalysisResponse(BaseModel):
    success: bool
    score: ATSScore
    insights: list[ATSInsight]
    recommendations: list[ATSRecommendation]
    matched_keywords: list[str]
    missing_keywords: list[str]
    experience_gaps: list[str]
    strengths: list[str]
    removable_words: list[str]  # Words/phrases that can be removed to improve ATS compatibility
    message: str

# New models for section optimization
class SectionOptimizationRequest(BaseModel):
    resume_data: ResumeData
    job_description: str
    section: str  # "summary", "experience", "skills", etc.
    section_data: dict  # The current section data
    custom_prompt: str = ""  # User's custom instructions

class SectionOptimizationResponse(BaseModel):
    success: bool
    optimized_section: dict
    explanation: str
    changes_made: list[str]
    message: str

# Add CORS middleware to allow frontend requests
app.add_middleware(
    CORSMiddleware,
    allow_origins=Config.CORS["allowed_origins"],
    allow_credentials=Config.CORS["allow_credentials"],
    allow_methods=Config.CORS["allow_methods"],
    allow_headers=Config.CORS["allow_headers"],
)

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file content."""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting text from PDF: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file content."""
    try:
        doc = Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting text from DOCX: {str(e)}")

def parse_resume_with_chatgpt(text: str) -> ResumeData:
    """Parse extracted text into structured resume data using ChatGPT."""
    try:
        # Check if OpenAI client is available
        if openai_client is None:
            raise Exception("OpenAI client not initialized")
        
        # Get OpenAI configuration for resume parsing
        openai_config = Config.get_openai_config("parsing")
        
        # Create the prompt for ChatGPT
        prompt = f"""
        Parse the following resume text and extract ALL available information into a JSON format. 
        Be thorough and accurate in your extraction. If information is not available, use empty strings or empty arrays.

        Required JSON structure:
        {{
            "name": "Full name of the person",
            "email": "Email address",
            "phone": "Phone number",
            "location": "City, State/Country (if mentioned)",
            "summary": "Professional summary or objective (create a concise summary from the content if not explicitly stated)",
            "github_profile": "GitHub profile URL or username",
            "linkedin_profile": "LinkedIn profile URL",
            "website": "Personal website or portfolio URL",
            "experience": [
                {{
                    "position": "Job title",
                    "company": "Company name",
                    "duration": "Employment duration (e.g., '2020-2023' or 'Jan 2020 - Present')",
                    "description": ["Bullet point 1", "Bullet point 2", "Bullet point 3", "Additional bullet points as needed"]
                }}
            ],
            "education": [
                {{
                    "degree": "Degree type and field",
                    "institution": "School/University name",
                    "year": "Graduation year",
                    "gpa": "GPA if mentioned",
                    "relevant_coursework": "Notable coursework if mentioned"
                }}
            ],
            "projects": [
                {{
                    "name": "Project name",
                    "description": ["Bullet point 1", "Bullet point 2", "Additional bullet points as needed"],
                    "technologies": "Technologies used",
                    "url": "Project URL if mentioned",
                    "duration": "Project duration if mentioned"
                }}
            ],
            "publications": [
                {{
                    "title": "Publication title",
                    "journal": "Journal or conference name",
                    "year": "Publication year",
                    "authors": "Authors (include the person if they are an author)",
                    "url": "Publication URL if mentioned"
                }}
            ],
            "certifications": [
                {{
                    "name": "Certification name",
                    "issuer": "Certifying organization",
                    "year": "Year obtained",
                    "expiry": "Expiry date if mentioned"
                }}
            ],
            "volunteer_experience": [
                {{
                    "position": "Volunteer position",
                    "organization": "Organization name",
                    "duration": "Duration of volunteer work",
                    "description": ["Bullet point 1", "Bullet point 2", "Additional bullet points as needed"]
                }}
            ],
            "awards": ["List of awards, honors, or recognitions"],
            "languages": ["List of languages and proficiency levels"],
            "references": [
                {{
                    "name": "Reference name",
                    "title": "Reference title/position",
                    "company": "Reference company",
                    "contact": "Contact information if provided"
                }}
            ],
            "skills": ["List of technical and professional skills"]
        }}

        Resume text:
        {text}

        IMPORTANT FORMATTING RULES:
        1. For experience descriptions: ALWAYS break down the job description into AT LEAST 3 separate bullet points as an array. Even if the original text is a single paragraph, analyze the content and create 3+ logical bullet points that capture different aspects of the role (e.g., responsibilities, achievements, technologies used, impact, etc.).
        2. For projects and volunteer_experience descriptions: Break down descriptions into arrays with at least 2-3 bullet points per project/volunteer role.
        3. Each bullet point should be a complete, meaningful statement that stands alone.
        4. Preserve the original meaning and content while making it more readable and structured.
        5. Use arrays for description fields instead of newline-separated strings.

        IMPORTANT: Extract ALL information that is present. Look for:
        - Social media profiles (GitHub, LinkedIn, Twitter, etc.)
        - Personal websites or portfolios
        - Location information
        - Projects with descriptions and technologies
        - Publications, papers, or research
        - Certifications and licenses
        - Volunteer work
        - Awards and honors
        - Languages spoken
        - References
        - Any other relevant professional information

        Return only the JSON object, no additional text or formatting.
        """

        # Call OpenAI API
        response = openai_client.chat.completions.create(
            model=openai_config["model"],
            messages=[
                {"role": "system", "content": "You are an expert resume parser. Extract information accurately and return only valid JSON."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=openai_config["max_tokens"],
            temperature=openai_config["temperature"]
        )

        # Extract the response content
        response_text = response.choices[0].message.content.strip()
        
        # Clean up the response (remove any markdown formatting)
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
        
        # Parse the JSON response
        parsed_data = json.loads(response_text)
        
        # Convert parsed data to structured models
        experience_list = []
        for exp in parsed_data.get("experience", []):
            experience_list.append(Experience(
                position=exp.get("position", ""),
                company=exp.get("company", ""),
                duration=exp.get("duration", ""),
                description=exp.get("description", [])
            ))
        
        education_list = []
        for edu in parsed_data.get("education", []):
            education_list.append(Education(
                degree=edu.get("degree", ""),
                institution=edu.get("institution", ""),
                year=edu.get("year", ""),
                gpa=edu.get("gpa", ""),
                relevant_coursework=edu.get("relevant_coursework", "")
            ))
        
        projects_list = []
        for proj in parsed_data.get("projects", []):
            projects_list.append(Project(
                name=proj.get("name", ""),
                description=proj.get("description", []),
                technologies=proj.get("technologies", ""),
                url=proj.get("url", ""),
                duration=proj.get("duration", "")
            ))
        
        publications_list = []
        for pub in parsed_data.get("publications", []):
            publications_list.append(Publication(
                title=pub.get("title", ""),
                journal=pub.get("journal", ""),
                year=pub.get("year", ""),
                authors=pub.get("authors", ""),
                url=pub.get("url", "")
            ))
        
        certifications_list = []
        for cert in parsed_data.get("certifications", []):
            certifications_list.append(Certification(
                name=cert.get("name", ""),
                issuer=cert.get("issuer", ""),
                year=cert.get("year", ""),
                expiry=cert.get("expiry", "")
            ))
        
        volunteer_list = []
        for vol in parsed_data.get("volunteer_experience", []):
            volunteer_list.append(VolunteerExperience(
                position=vol.get("position", ""),
                organization=vol.get("organization", ""),
                duration=vol.get("duration", ""),
                description=vol.get("description", [])
            ))
        
        references_list = []
        for ref in parsed_data.get("references", []):
            references_list.append(Reference(
                name=ref.get("name", ""),
                title=ref.get("title", ""),
                company=ref.get("company", ""),
                contact=ref.get("contact", "")
            ))
        
        # Validate and create ResumeData object
        return ResumeData(
            name=parsed_data.get("name", "Resume Owner"),
            email=parsed_data.get("email", ""),
            phone=parsed_data.get("phone", ""),
            location=parsed_data.get("location", ""),
            summary=parsed_data.get("summary", "Professional with relevant experience and skills."),
            github_profile=parsed_data.get("github_profile", ""),
            linkedin_profile=parsed_data.get("linkedin_profile", ""),
            website=parsed_data.get("website", ""),
            experience=experience_list,
            education=education_list,
            skills=parsed_data.get("skills", []),
            publications=publications_list,
            projects=projects_list,
            certifications=certifications_list,
            languages=parsed_data.get("languages", []),
            volunteer_experience=volunteer_list,
            awards=parsed_data.get("awards", []),
            references=references_list
        )
        
    except json.JSONDecodeError as e:
        print(f"JSON parsing error: {e}")
        raise HTTPException(status_code=500, detail="Failed to parse resume with ChatGPT. Please ensure your resume is well-formatted and try again.")
    except Exception as e:
        print(f"ChatGPT parsing error: {e}")
        raise HTTPException(status_code=500, detail="Failed to parse resume with ChatGPT. Please check your OpenAI API key and try again.")


def analyze_resume_with_ats(resume_data: ResumeData, job_description: str) -> ATSAnalysisResponse:
    """Analyze resume against job description using ChatGPT to simulate ATS analysis."""
    try:
        # Check if OpenAI client is available
        if openai_client is None:
            raise Exception("OpenAI client not initialized")
        
        # Get OpenAI configuration for resume optimization
        openai_config = Config.get_openai_config("optimization")
        
        # Create the prompt for ATS analysis
        prompt = f"""
        You are an expert ATS (Applicant Tracking System) analyst. Analyze the following resume against the job description and provide a comprehensive assessment similar to what a real ATS would generate.

        Resume Information:
        Name: {resume_data.name}
        Email: {resume_data.email}
        Location: {resume_data.location}
        Professional Summary: {resume_data.summary}
        
        Skills: {', '.join(resume_data.skills)}
        
        Experience:
        {chr(10).join([f"- {exp.position} at {exp.company} ({exp.duration}): {'; '.join(exp.description)}" for exp in resume_data.experience])}
        
        Education:
        {chr(10).join([f"- {edu.degree} from {edu.institution} ({edu.year})" for edu in resume_data.education])}
        
        Projects:
        {chr(10).join([f"- {proj.name}: {'; '.join(proj.description)} (Technologies: {proj.technologies})" for proj in resume_data.projects])}
        
        Certifications:
        {chr(10).join([f"- {cert.name} from {cert.issuer} ({cert.year})" for cert in resume_data.certifications])}
        
        Job Description:
        {job_description}

        IMPORTANT: Be strict and realistic in your analysis. Real ATS systems are unforgiving. Score conservatively and identify real weaknesses.

        Please provide a comprehensive ATS analysis in the following JSON format:

        {{
            "score": {{
                "overall_score": 85,
                "keyword_match_score": 78,
                "experience_relevance": 90,
                "education_fit": 88,
                "skills_alignment": 82
            }},
            "insights": [
                {{
                    "category": "strength",
                    "title": "Strong Technical Background",
                    "description": "Candidate demonstrates excellent technical skills relevant to the position",
                    "impact": "high"
                }},
                {{
                    "category": "weakness", 
                    "title": "Limited Industry Experience",
                    "description": "Candidate lacks specific experience in the target industry",
                    "impact": "medium"
                }}
            ],
            "recommendations": [
                {{
                    "title": "Add Industry-Specific Keywords",
                    "description": "Include more industry-specific terminology and technologies",
                    "priority": "high",
                    "effort": "easy"
                }}
            ],
            "matched_keywords": ["Python", "Machine Learning", "Data Analysis"],
            "missing_keywords": ["TensorFlow", "AWS", "Docker"],
            "experience_gaps": ["Cloud computing experience", "Team leadership"],
            "strengths": ["Strong technical skills", "Relevant education", "Project experience"],
            "removable_words": ["passionate about", "detail-oriented", "team player", "hard worker", "excellent communication skills"]
        }}

        STRICT Analysis Guidelines:
        1. Score conservatively (0-100) - real ATS systems are harsh
        2. Be unforgiving about missing required skills/qualifications
        3. Identify specific keywords that match and are missing from job description
        4. Highlight experience gaps that would cause immediate rejection
        5. Provide actionable recommendations with priority levels
        6. Focus on exact keyword matches - partial matches score lower
        7. Consider industry-specific requirements and certifications
        8. Identify filler words/phrases that add no value and should be removed
        9. Flag outdated or irrelevant skills/technologies
        10. Assess if candidate meets minimum requirements (education, experience years)
        11. Be critical of vague descriptions and generic statements
        12. Penalize heavily for missing must-have qualifications

        Return only the JSON object, no additional text or formatting.
        """

        # Call OpenAI API
        response = openai_client.chat.completions.create(
            model=openai_config["model"],
            messages=[
                {"role": "system", "content": "You are an expert ATS analyst. Provide detailed, actionable feedback in JSON format. Be specific and practical in your recommendations."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=openai_config["max_tokens"],
            temperature=openai_config["temperature"]
        )

        # Extract the response content
        response_text = response.choices[0].message.content.strip()
        
        # Clean up the response (remove any markdown formatting)
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
        
        # Parse the JSON response
        parsed_data = json.loads(response_text)
        
        # Create response objects
        score = ATSScore(
            overall_score=parsed_data["score"]["overall_score"],
            keyword_match_score=parsed_data["score"]["keyword_match_score"],
            experience_relevance=parsed_data["score"]["experience_relevance"],
            education_fit=parsed_data["score"]["education_fit"],
            skills_alignment=parsed_data["score"]["skills_alignment"]
        )
        
        insights = [
            ATSInsight(
                category=insight["category"],
                title=insight["title"],
                description=insight["description"],
                impact=insight["impact"]
            ) for insight in parsed_data["insights"]
        ]
        
        recommendations = [
            ATSRecommendation(
                title=rec["title"],
                description=rec["description"],
                priority=rec["priority"],
                effort=rec["effort"]
            ) for rec in parsed_data["recommendations"]
        ]
        
        return ATSAnalysisResponse(
            success=True,
            score=score,
            insights=insights,
            recommendations=recommendations,
            matched_keywords=parsed_data["matched_keywords"],
            missing_keywords=parsed_data["missing_keywords"],
            experience_gaps=parsed_data["experience_gaps"],
            strengths=parsed_data["strengths"],
            removable_words=parsed_data["removable_words"],
            message="ATS analysis completed successfully"
        )
        
    except json.JSONDecodeError as e:
        print(f"JSON parsing error: {e}")
        raise HTTPException(status_code=500, detail="Failed to parse ATS analysis response")
    except Exception as e:
        print(f"ATS analysis error: {e}")
        raise HTTPException(status_code=500, detail=f"Error during ATS analysis: {str(e)}")


def optimize_section_with_chatgpt(resume_data: ResumeData, job_description: str, section: str, section_data: dict, custom_prompt: str = "") -> SectionOptimizationResponse:
    """Optimize a specific resume section using ChatGPT with custom user instructions."""
    try:
        # Check if OpenAI client is available
        if openai_client is None:
            raise Exception("OpenAI client not initialized")
        
        # Get OpenAI configuration for optimization
        openai_config = Config.get_openai_config("optimization")
        
        # Create section-specific prompts
        section_prompts = {
            "summary": "professional summary or objective",
            "experience": "work experience entries",
            "skills": "skills list",
            "education": "education entries",
            "projects": "project entries",
            "publications": "publication entries",
            "certifications": "certification entries",
            "volunteer_experience": "volunteer experience entries",
            "awards": "awards and honors",
            "languages": "languages list",
            "references": "reference entries"
        }
        
        section_name = section_prompts.get(section, section)
        
        # Create the optimization prompt
        base_prompt = f"""
        You are an expert resume optimization specialist. Your task is to optimize the {section_name} section of a resume to better match a specific job description.

        JOB DESCRIPTION:
        {job_description}

        CURRENT RESUME CONTEXT:
        Name: {resume_data.name}
        Current Summary: {resume_data.summary}
        Skills: {', '.join(resume_data.skills)}
        
        CURRENT {section.upper()} SECTION DATA:
        {json.dumps(section_data, indent=2)}

        OPTIMIZATION GUIDELINES:
        1. Analyze the job description to identify key requirements, skills, and keywords
        2. Optimize the section to better align with the job requirements
        3. Maintain authenticity and truthfulness - only enhance what's already there
        4. Use action verbs and quantifiable achievements where possible
        5. Ensure the optimized content flows naturally and professionally
        6. Keep the same structure and format as the original
        7. Focus on relevance to the specific job posting

        IMPORTANT RULES:
        - Do NOT add false information or experiences
        - Do NOT change dates, company names, or other factual details
        - Do NOT make the content longer than necessary
        - Maintain the original tone and style
        - Only enhance and rephrase existing content for better impact
        """

        # Add custom user instructions if provided
        if custom_prompt.strip():
            base_prompt += f"""

        CUSTOM USER INSTRUCTIONS:
        {custom_prompt}

        Please incorporate these specific instructions while optimizing the section.
        """

        # Add section-specific instructions
        if section == "summary":
            base_prompt += """
        
        For the summary section:
        - Make it more targeted to the specific job
        - Highlight the most relevant skills and experiences
        - Keep it concise (2-3 sentences)
        - Use keywords from the job description naturally
        """
        elif section == "experience":
            base_prompt += """
        
        For experience entries:
        - Start bullet points with strong action verbs
        - Quantify achievements with numbers, percentages, or timeframes
        - Focus on results and impact rather than just responsibilities
        - Use keywords from the job description
        - Keep each bullet point concise but impactful
        """
        elif section == "skills":
            base_prompt += """
        
        For skills section:
        - Prioritize skills mentioned in the job description
        - Group related skills together
        - Use the exact terminology from the job posting
        - Remove outdated or irrelevant skills if space is limited
        - Add proficiency levels if appropriate
        """

        base_prompt += """

        Please return your response in the following JSON format:
        {
            "optimized_section": {
                // The optimized section data in the same structure as the input
            },
            "explanation": "Brief explanation of what was optimized and why",
            "changes_made": [
                "List of specific changes made",
                "Each change as a separate string"
            ]
        }

        Return only the JSON object, no additional text or formatting.
        """

        # Call OpenAI API
        response = openai_client.chat.completions.create(
            model=openai_config["model"],
            messages=[
                {"role": "system", "content": "You are an expert resume optimization specialist. Optimize resume sections to better match job requirements while maintaining authenticity and truthfulness."},
                {"role": "user", "content": base_prompt}
            ],
            max_tokens=openai_config["max_tokens"],
            temperature=openai_config["temperature"]
        )

        # Extract the response content
        response_text = response.choices[0].message.content.strip()
        
        # Clean up the response (remove any markdown formatting)
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
        
        # Parse the JSON response
        parsed_data = json.loads(response_text)
        
        return SectionOptimizationResponse(
            success=True,
            optimized_section=parsed_data["optimized_section"],
            explanation=parsed_data["explanation"],
            changes_made=parsed_data["changes_made"],
            message="Section optimization completed successfully"
        )
        
    except json.JSONDecodeError as e:
        print(f"JSON parsing error: {e}")
        raise HTTPException(status_code=500, detail="Failed to parse optimization response")
    except Exception as e:
        print(f"Section optimization error: {e}")
        raise HTTPException(status_code=500, detail=f"Error during section optimization: {str(e)}")


def generate_pdf(resume_data: ResumeData) -> bytes:
    """Generate PDF from resume data using ReportLab."""
    try:
        pdf_config = Config.PDF_GENERATION
        margins = pdf_config["margins"]
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=letter, 
            rightMargin=margins["right"],
            leftMargin=margins["left"],
            topMargin=margins["top"],
            bottomMargin=margins["bottom"]
        )
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Create custom styles
        fonts = pdf_config["fonts"]
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=20,
            alignment=1,  # Center alignment
            fontName='Helvetica-Bold'
        )
        
        contact_style = ParagraphStyle(
            'ContactStyle',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=25,
            alignment=1,  # Center alignment
            fontName='Helvetica'
        )
        
        header_style = ParagraphStyle(
            'CustomHeader',
            parent=styles['Heading2'],
            fontSize=12,
            spaceAfter=4,
            spaceBefore=15,
            fontName='Helvetica-Bold',
            textColor='#000000'
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=4,
            fontName='Helvetica'
        )
        
        company_style = ParagraphStyle(
            'CompanyStyle',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=2,
            fontName='Helvetica-Bold'
        )
        
        position_style = ParagraphStyle(
            'PositionStyle',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=4,
            leftIndent=20,
            fontName='Helvetica'
        )
        
        bullet_style = ParagraphStyle(
            'BulletStyle',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=2,
            leftIndent=20,
            fontName='Helvetica'
        )
        
        # Build content
        story = []
        
        # Title
        story.append(Paragraph(resume_data.name.upper(), title_style))
        
        # Contact information - single line format
        contact_parts = []
        if resume_data.location:
            contact_parts.append(resume_data.location)
        if resume_data.phone:
            contact_parts.append(resume_data.phone)
        if resume_data.email:
            contact_parts.append(resume_data.email)
        if resume_data.linkedin_profile:
            contact_parts.append(resume_data.linkedin_profile)
        
        if contact_parts:
            story.append(Paragraph(" • ".join(contact_parts), contact_style))
        
        # Professional Summary
        story.append(Paragraph("SUMMARY", header_style))
        story.append(HRFlowable(width="100%", thickness=1, lineCap='round', color=black, spaceAfter=4, spaceBefore=0))
        story.append(Paragraph(resume_data.summary, normal_style))
        
        # Skills & Other
        story.append(Paragraph("SKILLS & OTHER", header_style))
        story.append(HRFlowable(width="100%", thickness=1, lineCap='round', color=black, spaceAfter=4, spaceBefore=0))
        story.append(Paragraph(f"<b>Skills:</b> {', '.join(resume_data.skills)}", normal_style))
        
        # Add volunteering if exists
        if resume_data.volunteer_experience:
            vol_text = ", ".join([f"{vol.position} at {vol.organization} ({vol.duration})" for vol in resume_data.volunteer_experience])
            story.append(Paragraph(f"<b>Volunteering:</b> {vol_text}", normal_style))
        
        
        # Professional Experience
        story.append(Paragraph("PROFESSIONAL EXPERIENCE", header_style))
        story.append(HRFlowable(width="100%", thickness=1, lineCap='round', color=black, spaceAfter=4, spaceBefore=0))
        for exp in resume_data.experience:
            # Company, Location, Date format
            story.append(Paragraph(f"{exp.company}, {resume_data.location} ({exp.duration})", company_style))
            # Position indented
            story.append(Paragraph(exp.position, position_style))
            # Bullet points
            for bullet in exp.description:
                story.append(Paragraph(f"• {bullet}", bullet_style))
        
        # Education
        story.append(Paragraph("EDUCATION", header_style))
        story.append(HRFlowable(width="100%", thickness=1, lineCap='round', color=black, spaceAfter=4, spaceBefore=0))
        for edu in resume_data.education:
            # Institution, Location, Date format
            story.append(Paragraph(f"{edu.institution}, {resume_data.location} ({edu.year})", company_style))
            # Degree indented
            story.append(Paragraph(edu.degree, position_style))
            # Awards/details as bullets
            if edu.gpa or edu.relevant_coursework:
                if edu.gpa:
                    story.append(Paragraph(f"• Awards: {edu.gpa}", bullet_style))
                if edu.relevant_coursework:
                    story.append(Paragraph(f"• {edu.relevant_coursework}", bullet_style))
        
        
        # Projects (including Publications)
        if resume_data.projects or resume_data.publications:
            story.append(Paragraph("PROJECTS", header_style))
            story.append(HRFlowable(width="100%", thickness=1, lineCap='round', color=black, spaceAfter=4, spaceBefore=0))
            
            # Process projects and check for corresponding publications
            for project in resume_data.projects:
                # Project name, duration format
                story.append(Paragraph(f"{project.name} ({project.duration})", company_style))
                # Description as bullet points
                for bullet in project.description:
                    story.append(Paragraph(f"• {bullet}", bullet_style))
                
                # Check if there's a corresponding publication for this project
                corresponding_pub = None
                for pub in resume_data.publications:
                    if pub.title.lower() == project.name.lower():
                        corresponding_pub = pub
                        break
                
                # Add publication details if found
                if corresponding_pub:
                    if corresponding_pub.journal:
                        story.append(Paragraph(f"• Published in: {corresponding_pub.journal} ({corresponding_pub.year})", bullet_style))
                    if corresponding_pub.url:
                        story.append(Paragraph(f"• Publication URL: {corresponding_pub.url}", bullet_style))
            
            # Add standalone publications (those without corresponding projects)
            for pub in resume_data.publications:
                # Check if this publication already has a corresponding project
                has_corresponding_project = any(project.name.lower() == pub.title.lower() for project in resume_data.projects)
                
                if not has_corresponding_project:
                    # Publication title with "Publication" label
                    pub_text = f"{pub.title} (Publication)"
                    if pub.journal:
                        pub_text += f" - {pub.journal}"
                    if pub.year:
                        pub_text += f" ({pub.year})"
                    story.append(Paragraph(pub_text, company_style))
                    if pub.url:
                        story.append(Paragraph(f"• URL: {pub.url}", bullet_style))
        
        # Certifications
        if resume_data.certifications:
            story.append(Paragraph("CERTIFICATIONS", header_style))
            story.append(HRFlowable(width="100%", thickness=1, lineCap='round', color=black, spaceAfter=4, spaceBefore=0))
            for cert in resume_data.certifications:
                cert_text = f"<b>{cert.name}</b> - {cert.issuer}"
                if cert.year:
                    cert_text += f" ({cert.year})"
                if cert.expiry:
                    cert_text += f" | Expires: {cert.expiry}"
                story.append(Paragraph(cert_text, normal_style))
        
        
        # Awards
        if resume_data.awards:
            story.append(Paragraph("AWARDS & HONORS", header_style))
            story.append(HRFlowable(width="100%", thickness=1, lineCap='round', color=black, spaceAfter=4, spaceBefore=0))
            for award in resume_data.awards:
                story.append(Paragraph(f"• {award}", normal_style))
        
        # Languages
        if resume_data.languages:
            story.append(Paragraph("LANGUAGES", header_style))
            story.append(HRFlowable(width="100%", thickness=1, lineCap='round', color=black, spaceAfter=4, spaceBefore=0))
            story.append(Paragraph(" • ".join(resume_data.languages), normal_style))
        
        # References
        if resume_data.references:
            story.append(Paragraph("REFERENCES", header_style))
            story.append(HRFlowable(width="100%", thickness=1, lineCap='round', color=black, spaceAfter=4, spaceBefore=0))
            for ref in resume_data.references:
                ref_text = f"<b>{ref.name}</b>"
                if ref.title:
                    ref_text += f", {ref.title}"
                if ref.company:
                    ref_text += f" at {ref.company}"
                story.append(Paragraph(ref_text, normal_style))
                if ref.contact:
                    story.append(Paragraph(ref.contact, normal_style))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating PDF: {str(e)}")

def generate_docx(resume_data: ResumeData) -> bytes:
    """Generate DOCX from resume data using python-docx."""
    try:
        docx_config = Config.DOCX_GENERATION
        margins = docx_config["margins"]
        fonts = docx_config["fonts"]
        
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(margins["top"])
            section.bottom_margin = Inches(margins["bottom"])
            section.left_margin = Inches(margins["left"])
            section.right_margin = Inches(margins["right"])
        
        # Title
        title = doc.add_heading(resume_data.name.upper(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info
        contact_info = resume_data.email
        if resume_data.phone:
            contact_info += f" | {resume_data.phone}"
        if resume_data.location:
            contact_info += f" | {resume_data.location}"
        
        contact = doc.add_paragraph(contact_info)
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.runs[0].font.size = Pt(fonts["contact_size"])
        
        # Social profiles and website
        if resume_data.github_profile or resume_data.linkedin_profile or resume_data.website:
            social_info = []
            if resume_data.github_profile:
                social_info.append(f"GitHub: {resume_data.github_profile}")
            if resume_data.linkedin_profile:
                social_info.append(f"LinkedIn: {resume_data.linkedin_profile}")
            if resume_data.website:
                social_info.append(f"Website: {resume_data.website}")
            
            social = doc.add_paragraph(" • ".join(social_info))
            social.alignment = WD_ALIGN_PARAGRAPH.CENTER
            social.runs[0].font.size = Pt(fonts["contact_size"])
        
        # Add spacing
        doc.add_paragraph()
        
        # Professional Summary
        doc.add_heading('SUMMARY', level=1)
        summary_para = doc.add_paragraph(resume_data.summary)
        summary_para.runs[0].font.size = Pt(fonts["body_size"])
        
        # Skills & Other
        doc.add_heading('SKILLS & OTHER', level=1)
        skills_para = doc.add_paragraph(f"Skills: {', '.join(resume_data.skills)}")
        skills_para.runs[0].font.size = Pt(fonts["body_size"])
        
        # Add volunteering if exists
        if resume_data.volunteer_experience:
            vol_text = ", ".join([f"{vol.position} at {vol.organization} ({vol.duration})" for vol in resume_data.volunteer_experience])
            vol_para = doc.add_paragraph(f"Volunteering: {vol_text}")
            vol_para.runs[0].font.size = Pt(fonts["body_size"])
        
        
        # Professional Experience
        doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
        for exp in resume_data.experience:
            # Company, Location, Date format
            company_para = doc.add_paragraph()
            company_para.add_run(exp.company).bold = True
            company_para.add_run(f", {resume_data.location} ({exp.duration})")
            company_para.runs[0].font.size = Pt(fonts["body_size"])
            
            # Position indented
            position_para = doc.add_paragraph(f"    {exp.position}")
            position_para.runs[0].font.size = Pt(fonts["body_size"])
            
            # Job description bullets
            for bullet in exp.description:
                desc_para = doc.add_paragraph(f"    • {bullet}")
                desc_para.runs[0].font.size = Pt(fonts["body_size"])
            doc.add_paragraph()  # Add spacing
        
        # Education
        doc.add_heading('EDUCATION', level=1)
        for edu in resume_data.education:
            # Institution, Location, Date format
            institution_para = doc.add_paragraph()
            institution_para.add_run(edu.institution).bold = True
            institution_para.add_run(f", {resume_data.location} ({edu.year})")
            institution_para.runs[0].font.size = Pt(fonts["body_size"])
            
            # Degree indented
            degree_para = doc.add_paragraph(f"    {edu.degree}")
            degree_para.runs[0].font.size = Pt(fonts["body_size"])
            
            # Awards/details as bullets
            if edu.gpa:
                awards_para = doc.add_paragraph(f"    • Awards: {edu.gpa}")
                awards_para.runs[0].font.size = Pt(fonts["body_size"])
            if edu.relevant_coursework:
                coursework_para = doc.add_paragraph(f"    • {edu.relevant_coursework}")
                coursework_para.runs[0].font.size = Pt(fonts["body_size"])
        
        
        # Projects (including Publications)
        if resume_data.projects or resume_data.publications:
            doc.add_heading('PROJECTS', level=1)
            
            # Process projects and check for corresponding publications
            for project in resume_data.projects:
                # Project name, duration format
                project_para = doc.add_paragraph()
                project_text = project.name
                if project.duration:
                    project_text += f" ({project.duration})"
                project_para.add_run(project_text).bold = True
                project_para.runs[0].font.size = Pt(fonts["body_size"])
                
                # Description as bullet points
                for bullet in project.description:
                    bullet_para = doc.add_paragraph(f"    • {bullet}")
                    bullet_para.runs[0].font.size = Pt(fonts["body_size"])
                
                if project.url:
                    url_para = doc.add_paragraph(f"URL: {project.url}")
                    url_para.runs[0].font.size = Pt(fonts["body_size"])
                
                # Check if there's a corresponding publication for this project
                corresponding_pub = None
                for pub in resume_data.publications:
                    if pub.title.lower() == project.name.lower():
                        corresponding_pub = pub
                        break
                
                # Add publication details if found
                if corresponding_pub:
                    if corresponding_pub.journal:
                        journal_para = doc.add_paragraph(f"    • Published in: {corresponding_pub.journal} ({corresponding_pub.year})")
                        journal_para.runs[0].font.size = Pt(fonts["body_size"])
                    if corresponding_pub.url:
                        pub_url_para = doc.add_paragraph(f"    • Publication URL: {corresponding_pub.url}")
                        pub_url_para.runs[0].font.size = Pt(fonts["body_size"])
                
                doc.add_paragraph()  # Add spacing
            
            # Add standalone publications (those without corresponding projects)
            for pub in resume_data.publications:
                # Check if this publication already has a corresponding project
                has_corresponding_project = any(project.name.lower() == pub.title.lower() for project in resume_data.projects)
                
                if not has_corresponding_project:
                    pub_para = doc.add_paragraph()
                    pub_text = f"{pub.title} (Publication)"
                    if pub.journal:
                        pub_text += f" - {pub.journal}"
                    if pub.year:
                        pub_text += f" ({pub.year})"
                    pub_para.add_run(pub_text).bold = True
                    pub_para.runs[0].font.size = Pt(fonts["body_size"])
                    
                    if pub.url:
                        url_para = doc.add_paragraph(f"    • URL: {pub.url}")
                        url_para.runs[0].font.size = Pt(fonts["body_size"])
                    
                    doc.add_paragraph()  # Add spacing
        
        # Certifications
        if resume_data.certifications:
            doc.add_heading('CERTIFICATIONS', level=1)
            for cert in resume_data.certifications:
                cert_para = doc.add_paragraph()
                cert_text = f"{cert.name} - {cert.issuer}"
                if cert.year:
                    cert_text += f" ({cert.year})"
                if cert.expiry:
                    cert_text += f" | Expires: {cert.expiry}"
                cert_para.add_run(cert_text).bold = True
                cert_para.runs[0].font.size = Pt(fonts["body_size"])
        
        
        # Awards
        if resume_data.awards:
            doc.add_heading('AWARDS & HONORS', level=1)
            for award in resume_data.awards:
                award_para = doc.add_paragraph(f"• {award}")
                award_para.runs[0].font.size = Pt(fonts["body_size"])
        
        # Languages
        if resume_data.languages:
            doc.add_heading('LANGUAGES', level=1)
            lang_para = doc.add_paragraph(" • ".join(resume_data.languages))
            lang_para.runs[0].font.size = Pt(fonts["body_size"])
        
        # References
        if resume_data.references:
            doc.add_heading('REFERENCES', level=1)
            for ref in resume_data.references:
                ref_para = doc.add_paragraph()
                ref_text = ref.name
                if ref.title:
                    ref_text += f", {ref.title}"
                if ref.company:
                    ref_text += f" at {ref.company}"
                ref_para.add_run(ref_text).bold = True
                ref_para.runs[0].font.size = Pt(fonts["body_size"])
                
                if ref.contact:
                    contact_para = doc.add_paragraph(ref.contact)
                    contact_para.runs[0].font.size = Pt(fonts["body_size"])
                
                doc.add_paragraph()  # Add spacing
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating DOCX: {str(e)}")

@app.get("/")
async def root():
    """Health check endpoint."""
    return {"message": "Resume Text Extractor API is running"}

@app.post("/parse-resume", response_model=ParsedResumeResponse)
async def parse_resume(file: UploadFile = File(...)):
    """
    Parse uploaded resume file and return structured data.
    Supports PDF and DOCX formats.
    """
    # Validate file type
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    
    file_extension = file.filename.lower().split('.')[-1]
    
    if file_extension not in ['pdf', 'docx']:
        raise HTTPException(
            status_code=400, 
            detail="Unsupported file type. Please upload a PDF or DOCX file."
        )
    
    # Read file content
    try:
        file_content = await file.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading file: {str(e)}")
    
    # Extract text based on file type
    try:
        if file_extension == 'pdf':
            extracted_text = extract_text_from_pdf(file_content)
        elif file_extension == 'docx':
            extracted_text = extract_text_from_docx(file_content)
        
        # Parse the extracted text into structured data using ChatGPT
        if openai_client is None:
            raise HTTPException(
                status_code=503, 
                detail="Resume parsing service is currently unavailable. Please ensure OpenAI API key is configured and try again later."
            )
        
        parsed_data = parse_resume_with_chatgpt(extracted_text)
        
        return ParsedResumeResponse(
            success=True,
            data=parsed_data,
            message="Resume parsed successfully"
        )
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error parsing resume: {str(e)}")

@app.post("/generate-pdf")
async def generate_pdf_endpoint(resume_data: ResumeData):
    """
    Generate PDF from resume data.
    """
    try:
        pdf_content = generate_pdf(resume_data)
        
        return StreamingResponse(
            io.BytesIO(pdf_content),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={resume_data.name.replace(' ', '_')}_resume.pdf"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating PDF: {str(e)}")

@app.post("/generate-docx")
async def generate_docx_endpoint(resume_data: ResumeData):
    """
    Generate DOCX from resume data.
    """
    try:
        docx_content = generate_docx(resume_data)
        
        return StreamingResponse(
            io.BytesIO(docx_content),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={resume_data.name.replace(' ', '_')}_resume.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating DOCX: {str(e)}")

@app.post("/analyze-ats", response_model=ATSAnalysisResponse)
async def analyze_ats(request: ATSAnalysisRequest):
    """
    Analyze resume against job description using ChatGPT to simulate ATS analysis.
    Returns comprehensive scoring, insights, and recommendations.
    """
    try:
        if not request.job_description.strip():
            raise HTTPException(status_code=400, detail="Job description cannot be empty")
        
        if not request.resume_data:
            raise HTTPException(status_code=400, detail="Resume data is required")
        
        # Perform ATS analysis
        analysis_result = analyze_resume_with_ats(request.resume_data, request.job_description)
        
        return analysis_result
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error during ATS analysis: {str(e)}")

@app.post("/optimize-section", response_model=SectionOptimizationResponse)
async def optimize_section(request: SectionOptimizationRequest):
    """
    Optimize a specific resume section using ChatGPT with custom user instructions.
    Supports chatbot-style optimization with custom prompts.
    """
    try:
        if not request.job_description.strip():
            raise HTTPException(status_code=400, detail="Job description cannot be empty")
        
        if not request.resume_data:
            raise HTTPException(status_code=400, detail="Resume data is required")
        
        if not request.section.strip():
            raise HTTPException(status_code=400, detail="Section name is required")
        
        if not request.section_data:
            raise HTTPException(status_code=400, detail="Section data is required")
        
        # Validate section name
        valid_sections = [
            "summary", "experience", "skills", "education", "projects", 
            "publications", "certifications", "volunteer_experience", 
            "awards", "languages", "references"
        ]
        
        if request.section not in valid_sections:
            raise HTTPException(
                status_code=400, 
                detail=f"Invalid section. Must be one of: {', '.join(valid_sections)}"
            )
        
        # Perform section optimization
        optimization_result = optimize_section_with_chatgpt(
            request.resume_data,
            request.job_description,
            request.section,
            request.section_data,
            request.custom_prompt
        )
        
        return optimization_result
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error during section optimization: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    server_config = Config.SERVER
    uvicorn.run(
        app, 
        host=server_config["host"], 
        port=server_config["port"],
        debug=server_config["debug"]
    )
